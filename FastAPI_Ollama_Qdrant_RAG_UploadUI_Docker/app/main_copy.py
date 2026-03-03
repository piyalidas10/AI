import os
import shutil
import time
import uuid
import pandas as pd
from typing import List
from datetime import datetime

from fastapi import FastAPI, UploadFile, File, Form, Request, HTTPException
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates

from langchain_ollama import OllamaLLM, OllamaEmbeddings
from langchain_community.vectorstores import Qdrant
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains.retrieval import create_retrieval_chain

from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams
from pypdf import PdfReader
from docx import Document as DocxDocument


# =====================================================
# Configuration
# =====================================================

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://ollama:11434")
QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant:6333")

COLLECTION_NAME = "rag_collection"
UPLOAD_FOLDER = "uploaded_docs"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# =====================================================
# FastAPI App
# =====================================================

app = FastAPI(title="Enterprise RAG API")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
templates = Jinja2Templates(directory=os.path.join(BASE_DIR, "templates"))


# =====================================================
# LLM & Embeddings
# =====================================================

llm = OllamaLLM(
    model="phi3",
    base_url=OLLAMA_BASE_URL,
)

embeddings = OllamaEmbeddings(
    model="nomic-embed-text",
    base_url=OLLAMA_BASE_URL,
)


# =====================================================
# Globals
# =====================================================

vector_store = None
retrieval_chain = None
qdrant_client = None


# =====================================================
# STARTUP EVENT (Production Safe)
# =====================================================

@app.on_event("startup")
async def startup_event():
    global vector_store, retrieval_chain, qdrant_client

    print("🔄 Waiting for Qdrant to be ready...")

    # Retry connection (important in Docker)
    for i in range(10):
        try:
            qdrant_client = QdrantClient(url=QDRANT_URL)
            qdrant_client.get_collections()
            print("✅ Connected to Qdrant")
            break
        except Exception:
            print(f"⏳ Retry {i+1}/10 - Qdrant not ready...")
            time.sleep(3)
    else:
        raise Exception("❌ Could not connect to Qdrant")

    # Ensure Collection Exists
    # ✅ First time → collection exists → upload works
    # ❌ If deleted → upload tries to insert into non-existing collection → 500 Internal Server Error

    collections = qdrant_client.get_collections().collections
    collection_names = [c.name for c in collections]

    if COLLECTION_NAME not in collection_names:
        print("📦 Creating rag_collection...")
        qdrant_client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(
                size=768,  # nomic-embed-text dimension
                distance=Distance.COSINE,
            ),
        )
        print("✅ rag_collection created")
    else:
        print("✅ rag_collection already exists")

    # Initialize Vector Store
    vector_store = Qdrant(
        client=qdrant_client,
        collection_name=COLLECTION_NAME,
        embeddings=embeddings,
    )

    retriever = vector_store.as_retriever(search_kwargs={"k": 4})

    prompt = ChatPromptTemplate.from_template(
        """You are an assistant for question-answering tasks.
Use the retrieved context to answer the question.
If you don't know the answer, say you don't know.

Context:
{context}

Question:
{input}

Answer:"""
    )

    document_chain = create_stuff_documents_chain(llm, prompt)
    retrieval_chain = create_retrieval_chain(retriever, document_chain)

    print("🚀 Application startup complete")


# =====================================================
# Utility Functions
# =====================================================

def extract_text(file_path: str, filename: str, category: str) -> List[Document]:
    documents = []
    uploaded_time = datetime.utcnow().isoformat()

    if filename.lower().endswith(".pdf"):
        reader = PdfReader(file_path)

        for page_number, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            if page_text:
                documents.append(
                    Document(
                        page_content=page_text,
                        metadata={
                            "file_name": filename,
                            "page": page_number,
                            "uploaded_at": uploaded_time,
                            "category": category,
                        },
                    )
                )

    elif filename.lower().endswith(".docx"):
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "file_name": filename,
                    "page": 1,
                    "uploaded_at": uploaded_time,
                    "category": category,
                },
            )
        )

    elif filename.lower().endswith(".csv"):
        df = pd.read_csv(file_path)

        documents.append(
            Document(
                page_content=df.to_string(),
                metadata={
                    "file_name": filename,
                    "page": 1,
                    "uploaded_at": uploaded_time,
                    "category": category,
                },
            )
        )
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    return documents


def split_documents(documents: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
    )
    return splitter.split_documents(documents)


# =====================================================
# Routes
# =====================================================

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("upload.html", {"request": request})


@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    category: str = Form("general")
):
    if vector_store is None:
        raise HTTPException(status_code=503, detail="Vector store not ready")

    try:
        file_path = os.path.join(UPLOAD_FOLDER, file.filename)

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        documents = extract_text(file_path, file.filename, category)
        chunks = split_documents(documents)

        # Add UUID to avoid duplicate ID issue
        for chunk in chunks:
            chunk.metadata["doc_id"] = str(uuid.uuid4())

        vector_store.add_documents(chunks)

        return {
            "message": f"{file.filename} uploaded and indexed successfully",
            "category": category,
            "chunks_stored": len(chunks),
        }

    except Exception as e:
        print("❌ Upload Error:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/ask")
async def ask_question(
    question: str = Form(...),
    category: str = Form(None)
):
    if retrieval_chain is None:
        raise HTTPException(status_code=503, detail="QA chain not ready")

    try:
        if category:
            retriever = vector_store.as_retriever(
                search_kwargs={
                    "k": 4,
                    "filter": {
                        "must": [
                            {
                                "key": "category",
                                "match": {"value": category}
                            }
                        ]
                    }
                }
            )

            prompt = ChatPromptTemplate.from_template(
                """Use the context to answer.

Context:
{context}

Question:
{input}

Answer:"""
            )

            document_chain = create_stuff_documents_chain(llm, prompt)
            chain = create_retrieval_chain(retriever, document_chain)

            response = chain.invoke({"input": question})
        else:
            response = retrieval_chain.invoke({"input": question})

        return {
            "question": question,
            "answer": response["answer"],
        }

    except Exception as e:
        print("❌ QA Error:", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/health")
async def health():
    return {"status": "ok"}