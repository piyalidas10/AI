import os
import uuid
import shutil
from datetime import datetime
from typing import List

import pandas as pd
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.models.response_models import UploadResponse
from app.core.config import UPLOAD_FOLDER
from app.services.embedding_service import EmbeddingService

router = APIRouter(prefix="/documents", tags=["Documents"])

embedding_service = EmbeddingService()
embedding_service.connect()


# =====================================================
# Utility: Split Documents
# =====================================================

def split_documents(documents: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    return splitter.split_documents(documents)


# =====================================================
# Upload Endpoint
# =====================================================

@router.post("/upload", response_model=UploadResponse)
async def upload_file(
    file: UploadFile = File(...),
    category: str = Form("general")
):
    """
    Upload document and index into Qdrant vector store.
    Supported: PDF, DOCX, CSV
    """

    try:
        file_path = os.path.join(UPLOAD_FOLDER, file.filename)

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        documents = []
        uploaded_time = datetime.utcnow().isoformat()

        # ================= PDF =================
        if file.filename.lower().endswith(".pdf"):
            reader = PdfReader(file_path)

            for page_no, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                if text:
                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "file_name": file.filename,
                                "page": page_no,
                                "category": category,
                                "uploaded_at": uploaded_time,
                            },
                        )
                    )

        # ================= DOCX =================
        elif file.filename.lower().endswith(".docx"):
            doc = DocxDocument(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)

            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "file_name": file.filename,
                        "category": category,
                        "uploaded_at": uploaded_time,
                    },
                )
            )

        # ================= CSV =================
        elif file.filename.lower().endswith(".csv"):
            df = pd.read_csv(file_path)

            documents.append(
                Document(
                    page_content=df.to_string(),
                    metadata={
                        "file_name": file.filename,
                        "category": category,
                        "uploaded_at": uploaded_time,
                    },
                )
            )

        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type"
            )

        # ================= Chunking =================
        chunks = split_documents(documents)

        for chunk in chunks:
            chunk.metadata["doc_id"] = str(uuid.uuid4())

        # ================= Store in Qdrant =================
        embedding_service.ensure_collection()
        vector_store = embedding_service.get_vector_store()
        vector_store.add_documents(chunks)

        return UploadResponse(
            success=True,
            message="File uploaded and indexed successfully",
            chunks=len(chunks),
            file_name=file.filename
        )

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Upload error: {str(e)}"
        )