import os
import shutil
import uuid
import time
from typing import List

import numpy as np
import pandas as pd

from fastapi import FastAPI, UploadFile, File, Form, Request, HTTPException
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates

from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_ollama import OllamaLLM, OllamaEmbeddings
from langchain_qdrant import QdrantVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains.retrieval import create_retrieval_chain

from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams


# =====================================================
# CONFIG
# =====================================================

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://ollama:11434")
QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant:6333")

COLLECTION_NAME = "rag_collection"
UPLOAD_FOLDER = "uploaded_docs"
VECTOR_SIZE = 768

os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# =====================================================
# FASTAPI
# =====================================================

app = FastAPI(title="Enterprise RAG Platform")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
templates = Jinja2Templates(directory=os.path.join(BASE_DIR, "templates"))


# =====================================================
# GLOBAL OBJECTS
# =====================================================

qdrant_client: QdrantClient = None
vector_store = None
retrieval_chain = None

query_traces = []

# Store the latest metrics in memory for simplicity. In production, consider using a database or monitoring tool.
# Metrics include:
# - Faithfulness: How well the answer is supported by the retrieved context.
# - Answer Relevancy: How relevant the answer is to the question.
# - Context Precision: The proportion of retrieved context that is relevant to the answer.
# - Context Recall: The proportion of relevant context that was retrieved.
# - Answer Similarity: The average similarity between the answer and the retrieved context chunks.
# - Context Coverage: The proportion of the retrieved context that is covered by the answer.
# - Retrieval Score: The maximum similarity score between the answer and any retrieved chunk.
# - Hallucination Score: An estimate of how much the answer contains information not supported by the retrieved context.
# - Latency: The time taken to generate the answer.
# - Tokens: The total number of tokens in the question and answer.
latest_metrics = {
    "faithfulness": 0,
    "answer_relevancy": 0,
    "context_precision": 0,
    "context_recall": 0,
    "answer_similarity": 0,
    "context_coverage": 0,
    "retrieval_score": 0,
    "hallucination_score": 0,
    "latency": 0,
    "tokens": 0
}


# =====================================================
# COSINE SIMILARITY
# =====================================================

def cosine_similarity(a, b):

    a = np.array(a)
    b = np.array(b)

    return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b)))


# =====================================================
# EMBEDDINGS
# =====================================================

def get_embeddings():

    return OllamaEmbeddings(
        model="nomic-embed-text",
        base_url=OLLAMA_BASE_URL
    )


# =====================================================
# LLM
# =====================================================

def get_llm():

    return OllamaLLM(
        model="phi3",
        base_url=OLLAMA_BASE_URL,
        temperature=0
    )


# =====================================================
# VECTOR STORE
# =====================================================

def get_vector_store():

    global vector_store

    if vector_store is None:

        embeddings = get_embeddings()

        vector_store = QdrantVectorStore(
            client=qdrant_client,
            collection_name=COLLECTION_NAME,
            embedding=embeddings
        )

    return vector_store


# =====================================================
# CREATE COLLECTION
# =====================================================

def ensure_collection():

    collections = qdrant_client.get_collections().collections
    names = [c.name for c in collections]

    if COLLECTION_NAME not in names:

        qdrant_client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(
                size=VECTOR_SIZE,
                distance=Distance.COSINE
            )
        )


# =====================================================
# STARTUP
# =====================================================

@app.on_event("startup")
async def startup_event():

    global qdrant_client

    for i in range(10):

        try:

            qdrant_client = QdrantClient(url=QDRANT_URL)
            qdrant_client.get_collections()
            break

        except Exception:

            time.sleep(3)

    ensure_collection()


# =====================================================
# RAG CHAIN
# =====================================================

def get_rag_chain():

    global retrieval_chain

    if retrieval_chain is None:

        vs = get_vector_store()

        retriever = vs.as_retriever(
            search_type="mmr",
            search_kwargs={
                "k": 4,
                "fetch_k": 30,
                "lambda_mult": 0.6
            }
        )

        # MMR (Maximal Marginal Relevance) is a retrieval strategy that aims to balance relevance and diversity in the retrieved documents.
        # k: The number of documents to return.
        # fetch_k: The number of documents to fetch from the vector store before applying MMR.
        # lambda_mult: A parameter that controls the trade-off between relevance and diversity. A value

        # Small models like Phi-3 Mini, Gemma 2B, and Llama 3.2 3B need very strict prompts.
        # Use a grounded RAG prompt.

        prompt = ChatPromptTemplate.from_template(
            """
            Answer ONLY using the provided context.

            Context:
            {context}

            Question:
            {input}

            Answer:
            """
        )

        document_chain = create_stuff_documents_chain(
            get_llm(),
            prompt
        )

        retrieval_chain = create_retrieval_chain(
            retriever,
            document_chain
        )

    return retrieval_chain


# =====================================================
# TEXT EXTRACTION
# =====================================================

def extract_text(file_path: str, filename: str, category: str):

    documents = []

    if filename.endswith(".pdf"):

        reader = PdfReader(file_path)

        for i, page in enumerate(reader.pages):

            text = page.extract_text()

            if text:

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "file_name": filename,
                            "page": i + 1,
                            "category": category
                        }
                    )
                )

    elif filename.endswith(".docx"):

        doc = DocxDocument(file_path)

        text = "\n".join(p.text for p in doc.paragraphs)

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "file_name": filename,
                    "category": category
                }
            )
        )

    elif filename.endswith(".csv"):

        df = pd.read_csv(file_path)

        documents.append(
            Document(
                page_content=df.to_string(),
                metadata={
                    "file_name": filename,
                    "category": category
                }
            )
        )

    else:

        raise HTTPException(
            status_code=400,
            detail="Unsupported file format"
        )

    return documents


# =====================================================
# DOCUMENT SPLITTER
# =====================================================

def split_documents(documents: List[Document]):

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100
    )

    return splitter.split_documents(documents)


# =====================================================
# HEATMAP
# =====================================================

def compute_heatmap(question_emb, docs):

    embeddings = get_embeddings()

    heatmap = []

    for doc in docs:

        emb = embeddings.embed_query(doc.page_content)

        score = cosine_similarity(question_emb, emb)

        heatmap.append({
            "chunk": doc.page_content[:120],
            "score": round(score,3)
        })

    return heatmap


# =====================================================
# HALLUCINATION DETECTION
# =====================================================

def detect_hallucination(answer_emb, docs):

    embeddings = get_embeddings()

    scores = []

    for doc in docs:

        emb = embeddings.embed_query(doc.page_content)

        score = cosine_similarity(answer_emb, emb)

        scores.append(score)

    return round(1 - max(scores),3)


# =====================================================
# RAG EVALUATION
# =====================================================

def evaluate_rag(question, answer, docs, latency):

    embeddings = get_embeddings()

    q_emb = embeddings.embed_query(question)
    a_emb = embeddings.embed_query(answer)

    context = " ".join([d.page_content for d in docs])
    c_emb = embeddings.embed_query(context)

    faithfulness = cosine_similarity(a_emb, c_emb)
    answer_relevancy = cosine_similarity(q_emb, a_emb)
    context_recall = cosine_similarity(q_emb, c_emb)

    scores = []

    for doc in docs:

        emb = embeddings.embed_query(doc.page_content)

        score = cosine_similarity(a_emb, emb)

        scores.append(score)

    context_precision = sum([1 for s in scores if s > 0.5]) / len(scores)

    answer_similarity = np.mean(scores)
    context_coverage = sum(scores) / len(scores)
    retrieval_score = max(scores)

    hallucination_score = detect_hallucination(a_emb, docs)

    tokens = len(question.split()) + len(answer.split())

    return {
        "faithfulness": round(faithfulness,3),
        "answer_relevancy": round(answer_relevancy,3),
        "context_precision": round(context_precision,3),
        "context_recall": round(context_recall,3),
        "answer_similarity": round(answer_similarity,3),
        "context_coverage": round(context_coverage,3),
        "retrieval_score": round(retrieval_score,3),
        "hallucination_score": hallucination_score,
        "latency": round(latency,3),
        "tokens": tokens
    }


# =====================================================
# FILE UPLOAD
# =====================================================

@app.post("/upload")
async def upload_file(file: UploadFile = File(...), category: str = Form("general")):

    vs = get_vector_store()

    file_path = os.path.join(UPLOAD_FOLDER, file.filename)

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    documents = extract_text(file_path, file.filename, category)

    chunks = split_documents(documents)

    ids = [str(uuid.uuid4()) for _ in chunks]

    vs.add_documents(
        documents=chunks,
        ids=ids
    )

    return {"message": "File uploaded", "chunks": len(chunks)}


# =====================================================
# ASK QUESTION
# =====================================================

@app.post("/ask-ui", response_class=HTMLResponse)
async def ask_ui(request: Request, question: str = Form(...)):

    global latest_metrics, query_traces

    start = time.time()

    chain = get_rag_chain()
    vs = get_vector_store()

    docs = vs.similarity_search(question, k=4)

    response = chain.invoke({"input": question})

    answer = response["answer"]

    latency = time.time() - start

    latest_metrics = evaluate_rag(question, answer, docs, latency)

    embeddings = get_embeddings()
    q_emb = embeddings.embed_query(question)

    heatmap = compute_heatmap(q_emb, docs)

    trace = {
        "question": question,
        "answer": answer,
        "metrics": latest_metrics,
        "latency": latency,
        "heatmap": heatmap,
        "timestamp": time.time()
    }

    query_traces.append(trace)

    sources = []

    for doc in docs:

        sources.append({
            "file": doc.metadata.get("file_name"),
            "page": doc.metadata.get("page"),
            "category": doc.metadata.get("category")
        })

    return templates.TemplateResponse(
        "upload.html",
        {
            "request": request,
            "answer": answer,
            "sources": sources,
            "metrics": latest_metrics,
            "heatmap": heatmap
        }
    )


# =====================================================
# METRICS API
# =====================================================

@app.get("/rag-metrics")
def rag_metrics():

    return latest_metrics


# =====================================================
# DASHBOARD
# =====================================================

@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request):

    return templates.TemplateResponse(
        "dashboard.html",
        {
            "request": request,
            "metrics": latest_metrics
        }
    )


# =====================================================
# TRACES
# =====================================================

@app.get("/traces", response_class=HTMLResponse)
async def traces(request: Request):

    return templates.TemplateResponse(
        "traces.html",
        {
            "request": request,
            "traces": query_traces
        }
    )


# =====================================================
# HOME
# =====================================================

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):

    return templates.TemplateResponse(
        "upload.html",
        {"request": request}
    )


# =====================================================
# HEALTH
# =====================================================

@app.get("/health")
def health():

    return {"status": "ok"}