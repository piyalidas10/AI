import os
import shutil
import time
import uuid
import pandas as pd
from typing import List
from datetime import datetime
import json

from fastapi import FastAPI, UploadFile, File, Form, Request, HTTPException
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates

from langchain_ollama import OllamaLLM, OllamaEmbeddings
from langchain_community.vectorstores import Qdrant
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains.retrieval import create_retrieval_chain

from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams

from pypdf import PdfReader
from docx import Document as DocxDocument

# RAG Evaluation
from ragas import evaluate
# from ragas.metrics import faithfulness, answer_relevancy, context_precision
from ragas.metrics import faithfulness, answer_relevancy, context_utilization
from datasets import Dataset
from langchain_community.llms import Ollama

# Ragas internally tries to use OpenAI models via LangChain if you don't explicitly configure an LLM.
# This app uses Ollama (llama3), but RAGAS by default expects OpenAI API for evaluation.
# I must explicitly tell RAGAS to use your Ollama LLM instead of OpenAI.
from ragas.llms import LangchainLLMWrapper
from ragas.embeddings import LangchainEmbeddingsWrapper


# =====================================================
# CONFIG
# =====================================================

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://ollama:11434")
QDRANT_URL = os.getenv("QDRANT_URL", "http://qdrant:6333")

COLLECTION_NAME = "rag_collection"
VECTOR_SIZE = len(embeddings.embed_query("test"))
UPLOAD_FOLDER = "uploaded_docs"

RAG_LOG_FILE = "rag_logs.jsonl"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# =====================================================
# FastAPI App
# =====================================================

app = FastAPI(title="Enterprise RAG API")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
templates = Jinja2Templates(directory=os.path.join(BASE_DIR, "templates"))


# =====================================================
# LLM & Embeddings
# =====================================================

llm = OllamaLLM(
    model="llama3",
    base_url=OLLAMA_BASE_URL,
    temperature=0
)

embeddings = OllamaEmbeddings(
    model="nomic-embed-text",
    base_url=OLLAMA_BASE_URL
)

qdrant_client: QdrantClient = None
vector_store: Qdrant = None
retrieval_chain = None


# =====================================================
# COLLECTION ENSURE
# =====================================================

def ensure_collection_and_store():

    global vector_store

    collections = qdrant_client.get_collections().collections
    names = [c.name for c in collections]

    if COLLECTION_NAME not in names:
        print("📦 Creating rag_collection...")

        qdrant_client.create_collection(
            collection_name=COLLECTION_NAME,
            vectors_config=VectorParams(
                size=VECTOR_SIZE,
                distance=Distance.COSINE,
            ),
        )

        print("✅ rag_collection created")

    vector_store = Qdrant(
        client=qdrant_client,
        collection_name=COLLECTION_NAME,
        embeddings=embeddings,
    )


# =====================================================
# STARTUP
# =====================================================

@app.on_event("startup")
async def startup_event():

    global qdrant_client, retrieval_chain

    print("🔄 Connecting to Qdrant...")

    for i in range(10):

        try:
            qdrant_client = QdrantClient(url=QDRANT_URL)
            qdrant_client.get_collections()
            break

        except Exception:
            print(f"⏳ Retry {i+1}/10...")
            time.sleep(3)

    else:
        raise Exception("❌ Qdrant not reachable")

    ensure_collection_and_store()

    retriever = vector_store.as_retriever(
        search_kwargs={"k": 4}
    )

    prompt = ChatPromptTemplate.from_template(
        """
Use the context below to answer the question.

Context:
{context}

Question:
{input}

Answer clearly using only the context.
"""
    )

    document_chain = create_stuff_documents_chain(llm, prompt)

    retrieval_chain = create_retrieval_chain(
        retriever,
        document_chain
    )

    print("🚀 Startup complete")


# =====================================================
# TEXT EXTRACTION
# =====================================================

def extract_text(file_path: str, filename: str, category: str) -> List[Document]:

    documents = []
    uploaded_time = datetime.utcnow().isoformat()

    if filename.lower().endswith(".pdf"):

        reader = PdfReader(file_path)

        for page_no, page in enumerate(reader.pages, start=1):

            text = page.extract_text()

            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "file_name": filename,
                            "page": page_no,
                            "category": category,
                            "uploaded_at": uploaded_time,
                        },
                    )
                )

    elif filename.lower().endswith(".docx"):

        doc = DocxDocument(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "file_name": filename,
                    "category": category,
                    "uploaded_at": uploaded_time,
                },
            )
        )

    elif filename.lower().endswith(".csv"):

        df = pd.read_csv(file_path)

        documents.append(
            Document(
                page_content=df.to_string(),
                metadata={
                    "file_name": filename,
                    "category": category,
                    "uploaded_at": uploaded_time,
                },
            )
        )

    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    return documents


# =====================================================
# DOCUMENT SPLITTING
# =====================================================

def split_documents(documents: List[Document]) -> List[Document]:

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100
    )

    return splitter.split_documents(documents)


# =====================================================
# RAG EVALUATION
# =====================================================

def evaluate_rag(question, answer, contexts):

    # Ensure contexts format = List[List[str]]
    if isinstance(contexts, list):
        contexts_for_eval = [contexts]
    else:
        contexts_for_eval = [[contexts]]

    # ---------------------------------------------
    # Save RAG logs
    
    # Every question will generate a log entry like:
    #    {
    #    "timestamp": "2026-03-07T15:22:01",
    #    "question": "What are company office hours?",
    #    "answer": "Office hours are 9AM to 6PM.",
    #    "contexts": [
    #        "Office hours are 9AM to 6PM Monday-Friday",
    #        "Employees must log in before 9AM"
    #        ]
    #    }
    # ----------------------------------------------
    log_data = {
        "timestamp": datetime.utcnow().isoformat(),
        "question": question,
        "answer": answer,
        "contexts": contexts
    }

    with open(RAG_LOG_FILE, "a") as f:
        f.write(json.dumps(log_data) + "\n")

    # ----------------------------
    # Prepare dataset for RAGAS
    # ----------------------------
    data = {
        "question": [question],
        "answer": [answer],
        "contexts": contexts_for_eval
    }

    dataset = Dataset.from_dict(data)

    print("\nDEBUG DATASET:")
    print(dataset)

    # Wrap Ollama models for RAGAS
    # explicitly tell RAGAS to use your Ollama LLM instead of OpenAI because RAGAS by default expects OpenAI API for evaluation.
    ragas_llm = LangchainLLMWrapper(llm)
    ragas_embeddings = LangchainEmbeddingsWrapper(embeddings)

    # result = evaluate(dataset, metrics=[faithfulness, answer_relevancy, context_precision])
    result = evaluate(
        dataset,
        metrics=[
            faithfulness,
            answer_relevancy,
            context_utilization
        ],
        llm=ragas_llm,
        embeddings=ragas_embeddings,
        raise_exceptions=False # This will show the actual error instead of the generic message.
    )

    return result


# =====================================================
# ROUTES
# =====================================================

@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    category: str = Form("general")
):

    try:

        ensure_collection_and_store()

        file_path = os.path.join(UPLOAD_FOLDER, file.filename)

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        documents = extract_text(file_path, file.filename, category)

        chunks = split_documents(documents)

        for chunk in chunks:
            chunk.metadata["doc_id"] = str(uuid.uuid4())

        vector_store.add_documents(chunks)

        return {
            "message": "File uploaded successfully",
            "chunks": len(chunks)
        }

    except Exception as e:

        print("❌ Upload Error:", str(e))

        raise HTTPException(status_code=500, detail=str(e))


# =====================================================
# ASK QUESTION
# =====================================================

@app.post("/ask-ui", response_class=HTMLResponse)
async def ask_ui(request: Request, question: str = Form(...)):

    try:

        response = retrieval_chain.invoke(
            {"input": question}
        )

        answer = response["answer"]

        # Context Captured for Evaluation
        # Without this RAG evaluation does not work.
        # response["context"] comes from LangChain retriever.

        # Example retrieved documents:
        # response["context"] = [
        #    Document(page_content="Office hours 9AM-6PM"),
        #    Document(page_content="Company works Monday-Friday"),
        #    Document(page_content="Employees must login before 9AM")
        #]
        contexts = [
            doc.page_content
            for doc in response["context"]
        ]

        # contexts = [
        # "Office hours 9AM-6PM",
        # "Company works Monday-Friday",
        # "Employees must login before 9AM"
        # ]

        # Debug retrieved documents
        print("\n📄 Retrieved Documents:")
        for doc in response["context"]:
            print(doc.metadata)

        evaluation = evaluate_rag(
            question,
            answer,
            contexts
        )

        sources = [
            {
                "file_name": doc.metadata.get("file_name"),
                "page": doc.metadata.get("page"),
                "category": doc.metadata.get("category")
            }
            for doc in response["context"]
        ]

        return templates.TemplateResponse(
            "upload.html",
            {
                "request": request,
                "answer": answer,
                "evaluation": evaluation,
                "sources": sources
            }
        )

    except Exception as e:

        return templates.TemplateResponse(
            "upload.html",
            {
                "request": request,
                "answer": f"Error: {str(e)}"
            }
        )


# =====================================================
# HOME
# =====================================================

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):

    return templates.TemplateResponse(
        "upload.html",
        {"request": request}
    )


# =====================================================
# HEALTH
# =====================================================

@app.get("/health")
async def health():

    return {"status": "ok"}